from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from ..document import build_turns
from ..format_utils import fmt_hhmmss


def export_docx(segments, path: Path, title: str = "전사 결과") -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for t in build_turns(segments):
        speaker_label = t.speaker or "화자"
        heading = doc.add_paragraph()
        run = heading.add_run(f"{speaker_label}  [{fmt_hhmmss(t.start)} - {fmt_hhmmss(t.end)}]")
        run.bold = True
        doc.add_paragraph(t.text)

    doc.save(str(path))
